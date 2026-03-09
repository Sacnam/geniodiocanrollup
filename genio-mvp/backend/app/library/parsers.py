"""
Document Parsers for Library module.
From LIBRARY_PRD.md §3.1

Supports: EPUB, PDF, DOCX, HTML, Markdown
"""
import re
from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Tuple


class DocumentParser(ABC):
    """Base class for document parsers."""
    
    @abstractmethod
    def parse(self, file_path: str) -> Dict:
        """
        Parse document and return structured content.
        
        Returns:
            {
                "title": str,
                "author": str,
                "content": str,  # Markdown format
                "chapters": [{"title": str, "content": str}],
                "metadata": dict,
            }
        """
        pass
    
    def extract_chapters(self, content: str, chapter_patterns: List[str]) -> List[Tuple[str, str]]:
        """Extract chapters based on heading patterns."""
        chapters = []
        
        # Default patterns for chapter detection
        if not chapter_patterns:
            chapter_patterns = [
                r'#{1,2}\s+',  # Markdown headers
                r'^Chapter\s+\d+',  # Chapter X
                r'^\d+\.\s+',  # 1. Title
            ]
        
        # Split content by chapter markers
        current_chapter = "Introduction"
        current_content = []
        
        for line in content.split('\n'):
            is_chapter_start = False
            
            for pattern in chapter_patterns:
                if re.match(pattern, line, re.IGNORECASE):
                    # Save previous chapter
                    if current_content:
                        chapters.append((
                            current_chapter,
                            '\n'.join(current_content).strip()
                        ))
                    
                    # Start new chapter
                    current_chapter = line.strip('#').strip()
                    current_content = []
                    is_chapter_start = True
                    break
            
            if not is_chapter_start:
                current_content.append(line)
        
        # Don't forget last chapter
        if current_content:
            chapters.append((
                current_chapter,
                '\n'.join(current_content).strip()
            ))
        
        return chapters if chapters else [("Content", content)]


class EPUBParser(DocumentParser):
    """EPUB parser using ebooklib + BeautifulSoup."""
    
    def parse(self, file_path: str) -> Dict:
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("Install ebooklib and beautifulsoup4: pip install ebooklib beautifulsoup4")
        
        book = epub.read_epub(file_path)
        
        # Extract metadata
        title = book.get_metadata('DC', 'title')
        title = title[0][0] if title else "Untitled"
        
        author = book.get_metadata('DC', 'creator')
        author = author[0][0] if author else None
        
        # Extract content from chapters
        chapters = []
        full_content = []
        
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                
                # Get chapter title from first h1/h2
                heading = soup.find(['h1', 'h2'])
                chapter_title = heading.get_text() if heading else item.get_name()
                
                # Get text content
                text = soup.get_text()
                text = self._clean_text(text)
                
                if text.strip():
                    chapters.append({
                        "title": chapter_title,
                        "content": text
                    })
                    full_content.append(text)
        
        return {
            "title": title,
            "author": author,
            "content": "\n\n".join(full_content),
            "chapters": chapters,
            "metadata": {
                "format": "epub",
                "chapter_count": len(chapters),
            }
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()


class DOCXParser(DocumentParser):
    """DOCX parser using python-docx."""
    
    def parse(self, file_path: str) -> Dict:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
        
        doc = Document(file_path)
        
        # Extract metadata
        title = doc.core_properties.title or "Untitled"
        author = doc.core_properties.author
        
        # Extract paragraphs
        paragraphs = []
        chapters = []
        current_chapter_title = "Content"
        current_chapter_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect chapter headings by style
            if para.style.name.startswith('Heading'):
                # Save previous chapter
                if current_chapter_content:
                    chapters.append({
                        "title": current_chapter_title,
                        "content": '\n'.join(current_chapter_content)
                    })
                
                current_chapter_title = text
                current_chapter_content = []
            else:
                current_chapter_content.append(text)
            
            paragraphs.append(text)
        
        # Save last chapter
        if current_chapter_content:
            chapters.append({
                "title": current_chapter_title,
                "content": '\n'.join(current_chapter_content)
            })
        
        return {
            "title": title,
            "author": author,
            "content": '\n\n'.join(paragraphs),
            "chapters": chapters if len(chapters) > 1 else [{"title": "Content", "content": '\n\n'.join(paragraphs)}],
            "metadata": {
                "format": "docx",
                "chapter_count": len(chapters),
            }
        }


class HTMLParser(DocumentParser):
    """HTML parser using Readability-lxml."""
    
    def parse(self, file_path: str) -> Dict:
        try:
            from readability import Document as ReadabilityDocument
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("Install readability-lxml: pip install readability-lxml")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html = f.read()
        
        # Use Readability to extract main content
        doc = ReadabilityDocument(html)
        title = doc.title()
        summary = doc.summary()
        
        # Parse with BeautifulSoup
        soup = BeautifulSoup(summary, 'html.parser')
        
        # Extract text
        content = soup.get_text()
        content = self._clean_text(content)
        
        # Try to extract author
        author = None
        soup_full = BeautifulSoup(html, 'html.parser')
        
        # Look for author meta tags
        author_meta = soup_full.find('meta', attrs={'name': 'author'})
        if author_meta:
            author = author_meta.get('content')
        
        # Extract headings as chapters
        headings = soup.find_all(['h1', 'h2', 'h3'])
        chapters = []
        
        if headings:
            for h in headings:
                chapter_title = h.get_text().strip()
                # Get content until next heading
                chapter_content = []
                sibling = h.find_next_sibling()
                while sibling and sibling.name not in ['h1', 'h2', 'h3']:
                    if sibling.get_text().strip():
                        chapter_content.append(sibling.get_text().strip())
                    sibling = sibling.find_next_sibling()
                
                if chapter_content:
                    chapters.append({
                        "title": chapter_title,
                        "content": '\n\n'.join(chapter_content)
                    })
        
        if not chapters:
            chapters = [{"title": "Content", "content": content}]
        
        return {
            "title": title,
            "author": author,
            "content": content,
            "chapters": chapters,
            "metadata": {
                "format": "html",
                "source_url": None,  # Could be extracted from meta
            }
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()


class MarkdownParser(DocumentParser):
    """Markdown parser - native support."""
    
    def parse(self, file_path: str) -> Dict:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract title from first h1
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else "Untitled"
        
        # Try to extract YAML frontmatter
        author = None
        frontmatter_match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, re.DOTALL)
        if frontmatter_match:
            frontmatter = frontmatter_match.group(1)
            # Parse simple YAML
            for line in frontmatter.split('\n'):
                if line.startswith('author:'):
                    author = line.split(':', 1)[1].strip()
                elif line.startswith('title:'):
                    title = line.split(':', 1)[1].strip().strip('"\'')
        
        # Extract chapters by headings
        chapters = []
        current_title = "Introduction"
        current_content = []
        
        for line in content.split('\n'):
            if line.startswith('#'):
                # Save previous chapter
                if current_content:
                    chapters.append({
                        "title": current_title,
                        "content": '\n'.join(current_content).strip()
                    })
                current_title = line.lstrip('#').strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Save last chapter
        if current_content:
            chapters.append({
                "title": current_title,
                "content": '\n'.join(current_content).strip()
            })
        
        return {
            "title": title,
            "author": author,
            "content": content,
            "chapters": chapters if len(chapters) > 1 else [{"title": "Content", "content": content}],
            "metadata": {
                "format": "markdown",
                "chapter_count": len(chapters),
            }
        }


# Parser factory
PARSERS = {
    'epub': EPUBParser,
    'docx': DOCXParser,
    'html': HTMLParser,
    'markdown': MarkdownParser,
}


def get_parser(doc_type: str) -> DocumentParser:
    """Get appropriate parser for document type."""
    parser_class = PARSERS.get(doc_type.lower())
    if not parser_class:
        raise ValueError(f"Unsupported document type: {doc_type}")
    return parser_class()


def parse_document(file_path: str, doc_type: str) -> Dict:
    """
    Parse document using appropriate parser.
    
    Args:
        file_path: Path to document file
        doc_type: Document type (epub, docx, html, markdown)
    
    Returns:
        Parsed document structure
    """
    parser = get_parser(doc_type)
    return parser.parse(file_path)
